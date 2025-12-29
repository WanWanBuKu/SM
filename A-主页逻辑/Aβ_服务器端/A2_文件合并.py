#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件名: file_merger.py
功能: 服务器端文件合并工具
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime
import json

# 尝试导入Word处理库
try:
    from docx import Document
    from docxcompose.composer import Composer
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class FileMerger:
    """文件合并器"""
    
    def __init__(self):
        self.default_dir = "/storage/emulated/0/Download/"
        self.supported_types = {
            "1": {"name": "TXT文件", "pattern": "*.txt", "output": "合并_TXT文件.txt"},
            "2": {"name": "Python文件", "pattern": "*.py", "output": "合并_Python文件.txt"},
            "3": {"name": "DOCX文件", "pattern": "*.docx", "output": "合并_DOCX文件.docx"},
            "4": {"name": "除HTML外的所有文件", "pattern": "*", "output": "合并_除HTML外所有文件.txt"}
        }
    
    def natural_sort_key(self, filename):
        """自然排序函数"""
        return [int(text) if text.isdigit() else text.lower() 
                for text in re.split(r'(\d+)', filename)]
    
    def get_files_by_type(self, directory, merge_type):
        """根据合并类型获取文件列表"""
        if merge_type not in self.supported_types:
            return [], "", ""
        
        type_info = self.supported_types[merge_type]
        pattern = type_info["pattern"]
        output_name = type_info["output"]
        merge_type_name = type_info["name"]
        
        # 获取文件列表
        if merge_type == "4":  # 除HTML外的所有文件
            all_files = list(Path(directory).glob("*"))
            files = [f for f in all_files 
                    if f.is_file() 
                    and f.suffix.lower() not in ['.html', '.htm']
                    and f.name != output_name]
        else:
            files = list(Path(directory).glob(pattern))
        
        # 排除输出文件
        files = [f for f in files if f.name != output_name]
        
        # 自然排序
        files.sort(key=lambda x: self.natural_sort_key(x.name))
        
        return files, output_name, merge_type_name
    
    def merge_text_files(self, files, output_path, merge_type_name, directory):
        """合并文本文件（TXT、Python、其他）"""
        try:
            with open(output_path, 'w', encoding='utf-8') as outfile:
                # 写入文件大纲
                outfile.write("===== 文件大纲 =====\n")
                outfile.write(f"合并时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                outfile.write(f"合并类型: {merge_type_name}\n")
                outfile.write(f"文件数量: {len(files)}\n")
                outfile.write(f"目录路径: {directory}\n")
                outfile.write("=========================\n\n")
                
                # 列出所有文件
                for i, file_path in enumerate(files, 1):
                    outfile.write(f"{i:2d}. {file_path.name}\n")
                
                outfile.write("\n===== 文件内容 =====\n\n")
                
                # 写入每个文件的内容
                for file_path in files:
                    outfile.write(f"--- {file_path.name} ⬇️⬇️⬇️⬇️⬇️ ---\n")
                    
                    try:
                        file_size = file_path.stat().st_size
                        mod_time = datetime.fromtimestamp(file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                        
                        outfile.write(f"文件大小: {file_size} 字节\n")
                        outfile.write(f"修改时间: {mod_time}\n")
                        outfile.write("-------------------------\n")
                        
                        # 读取文件内容
                        try:
                            with open(file_path, 'r', encoding='utf-8') as infile:
                                content = infile.read()
                                outfile.write(content)
                        except UnicodeDecodeError:
                            try:
                                with open(file_path, 'r', encoding='gbk') as infile:
                                    content = infile.read()
                                    outfile.write(content)
                            except:
                                outfile.write("【无法读取文件内容：编码问题】\n")
                        except Exception as e:
                            outfile.write(f"【读取文件出错: {str(e)}】\n")
                            
                    except Exception as e:
                        outfile.write(f"【获取文件信息出错: {str(e)}】\n")
                    
                    outfile.write("\n")
            
            return True
            
        except Exception as e:
            return False
    
    def merge_docx_files(self, files, output_path):
        """合并DOCX文件"""
        if not DOCX_AVAILABLE:
            return False
        
        try:
            # 创建主文档
            master_doc = Document()
            composer = Composer(master_doc)
            
            # 创建目录文档
            toc_doc = Document()
            
            # 添加目录标题
            title_paragraph = toc_doc.add_paragraph()
            title_run = title_paragraph.add_run("===== 大纲 =====")
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            toc_doc.add_paragraph()  # 空行
            
            # 添加目录内容
            for i, file_path in enumerate(files, 1):
                clean_name = file_path.stem
                para = toc_doc.add_paragraph(f"{i:2d}. {clean_name}", style='ListNumber')
            
            # 合并目录页
            composer.append(toc_doc)
            
            # 合并所有文档内容
            for i, file_path in enumerate(files, 1):
                try:
                    clean_name = file_path.stem
                    
                    # 添加分隔符
                    separator_doc = Document()
                    self.add_file_separator(separator_doc, clean_name)
                    composer.append(separator_doc)
                    
                    # 合并原文档
                    source_doc = Document(str(file_path))
                    composer.append(source_doc)
                    
                except Exception as e:
                    continue
            
            # 保存文档
            composer.save(output_path)
            return True
            
        except Exception as e:
            return False
    
    def add_file_separator(self, doc, filename):
        """为DOCX文档添加文件分隔符"""
        doc.add_page_break()
        
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(f"--- {filename} ⬇️⬇️⬇️⬇️⬇️ ---")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def merge_files(self, merge_type, directory):
        """执行合并操作"""
        # 使用默认目录如果未提供
        if not directory:
            directory = self.default_dir
        
        # 验证目录存在
        if not os.path.exists(directory):
            return {"status": "error", "message": f"目录不存在: {directory}"}
        
        # 获取文件列表
        files, output_name, merge_type_name = self.get_files_by_type(directory, merge_type)
        
        if not files:
            return {"status": "error", "message": "未找到任何符合条件的文件"}
        
        # 设置输出文件路径
        output_path = os.path.join(directory, output_name)
        
        # 根据选择执行不同的合并
        if merge_type == "3":  # DOCX文件
            if not DOCX_AVAILABLE:
                return {"status": "error", "message": "DOCX合并功能不可用，请安装库：pip install python-docx docxcompose"}
            success = self.merge_docx_files(files, output_path)
        else:  # 文本文件
            success = self.merge_text_files(files, output_path, merge_type_name, directory)
        
        # 返回结果
        if success:
            output_size = os.path.getsize(output_path)
            return {
                "status": "success",
                "message": "文件合并完成",
                "output_file": output_path,
                "file_count": len(files),
                "file_size": output_size,
                "files": [f.name for f in files]
            }
        else:
            return {"status": "error", "message": "合并失败"}

    @staticmethod
    def get_page():
        """返回合并页面"""
        from flask import send_from_directory
        return send_from_directory('.', 'merge.html')
    
    @staticmethod
    def run_merge():
        """执行合并操作的路由处理函数"""
        from flask import request, jsonify
        try:
            data = request.get_json()
            merge_type = data.get('merge_type', '')
            directory = data.get('directory', '')
            
            if not merge_type:
                return jsonify({'status': 'error', 'message': '未提供合并类型'}), 400
            
            merger = FileMerger()
            result = merger.merge_files(merge_type, directory)
            
            if result['status'] == 'success':
                return jsonify(result), 200
            else:
                return jsonify(result), 400
                
        except Exception as e:
            return jsonify({'status': 'error', 'message': f'内部错误: {str(e)}'}), 500
